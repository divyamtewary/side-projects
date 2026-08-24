from __future__ import annotations
import pathlib
from typing import Dict, Any

def export_pdf(run_dir: pathlib.Path, output_path: pathlib.Path):
    """
    Minimal PDF export without heavy deps: uses reportlab if available, else fallback to writing overview.md renamed.
    Generates a simple PDF with summary and chart images if reportlab available.
    """
    overview = (run_dir / "overview.md")
    text = overview.read_text(encoding="utf-8") if overview.exists() else "No overview"
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak, HRFlowable, Table, TableStyle
        from reportlab.lib import colors

        doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                                leftMargin=18*mm, rightMargin=18*mm, topMargin=14*mm, bottomMargin=14*mm,
                                title=f"SLM Evaluation – {run_dir.name}", author="slm-evaluation-suite")

        styles = getSampleStyleSheet()
        h1 = ParagraphStyle('h1', parent=styles['Heading1'], fontSize=14, textColor=HexColor("#111827"), spaceAfter=6, spaceBefore=12)
        h2 = ParagraphStyle('h2', parent=styles['Heading2'], fontSize=10, textColor=HexColor("#4c1d95"), spaceAfter=4, spaceBefore=10)
        body = ParagraphStyle('body', parent=styles['BodyText'], fontSize=8, leading=11, textColor=HexColor("#1f2937"), spaceAfter=6)
        small = ParagraphStyle('small', parent=styles['BodyText'], fontSize=7, leading=9, textColor=HexColor("#6b7280"))
        caption = ParagraphStyle('cap', parent=styles['BodyText'], fontSize=7, leading=9, textColor=HexColor("#6b7280"), alignment=TA_CENTER)

        story = []
        # Title
        manifest = {}
        try:
            import json
            manifest = json.loads((run_dir / "manifest.json").read_text(encoding="utf-8"))
        except Exception:
            pass
        story.append(Paragraph(f"SLM Evaluation Suite — Report", ParagraphStyle('titlep', parent=styles['Title'], fontSize=16, textColor=HexColor("#111827"), alignment=TA_CENTER)))
        story.append(Paragraph(f"{manifest.get('model_id','Unknown model')}  •  {manifest.get('preset','')}  •  {manifest.get('created_utc','')}", ParagraphStyle('sub', parent=styles['Normal'], fontSize=8, textColor=HexColor("#6b7280"), alignment=TA_CENTER, spaceAfter=8)))
        story.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#e5e7eb"), spaceAfter=8))
        story.append(Paragraph(f"Run ID <font color=\"#6b7280\">{manifest.get('run_id','')}</font>  •  Schema {manifest.get('schema_version','1.0')}", small))
        story.append(Spacer(1, 6))

        # Render markdown as paragraphs (simplified)
        for line in text.split("\n"):
            line=line.strip()
            if not line:
                story.append(Spacer(1,4))
                continue
            if line.startswith("# "):
                story.append(Paragraph(line[2:].replace("`",""), h1))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:].replace("`",""), h2))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:].replace("`",""), h2))
            elif line.startswith("|") and "|" in line:
                # skip tables for now – render as preformatted small
                cleaned = line.replace("|"," | ")
                story.append(Paragraph(f"<font face=\"Courier\" size=6>{cleaned}</font>", small))
            elif line.startswith(">"):
                story.append(Paragraph(f"<i>{line[1:].strip()}</i>", ParagraphStyle('quote', parent=body, leftIndent=10, textColor=HexColor("#4b5563"), borderPadding=(4,4,4))))
            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph(f"• {line[2:]}", body))
            elif line.startswith("---"):
                story.append(HRFlowable(width="100%", thickness=0.4, color=HexColor("#e5e7eb"), spaceAfter=6))
            else:
                # escape < > for reportlab
                esc = line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                story.append(Paragraph(esc, body))

        story.append(Spacer(1, 10))
        story.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#e5e7eb")))
        story.append(Paragraph("Charts", h1))

        # Add chart images
        charts = sorted((run_dir / "charts").glob("*.png"))
        for ch in charts:
            try:
                story.append(Paragraph(ch.name, caption))
                img = Image(str(ch), width=160*mm, height=65*mm)
                img.hAlign = 'CENTER'
                story.append(img)
                story.append(Spacer(1, 6))
            except Exception as e:
                story.append(Paragraph(f"Chart {ch.name} could not be embedded: {e}", small))

        story.append(Spacer(1, 8))
        story.append(Paragraph("Generated by slm-evaluation-suite v0.1.0 — evidence cards, not a single quality score. Descriptive only.", small))

        doc.build(story)
        return output_path

    except ImportError as e:
        # fallback: copy markdown as .pdf.txt placeholder and also write simple text PDF via matplotlib
        fallback = output_path.with_suffix(".txt")
        fallback.write_text(text, encoding="utf-8")
        # try minimal PDF via matplotlib
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            fig = plt.figure(figsize=(8.27,11.69))
            fig.text(0.06,0.96,f"SLM Evaluation – {run_dir.name}", fontsize=12, weight="bold")
            # wrap text
            import textwrap
            wrapped = "\n".join(textwrap.wrap(text, width=110))
            fig.text(0.06,0.92,wrapped[:9000], fontsize=6, va="top", family="monospace")
            fig.text(0.06,0.03,"(Fallback PDF – install reportlab for richer layout: pip install reportlab)", fontsize=7, style="italic", color="gray")
            fig.savefig(str(output_path))
            plt.close(fig)
            return output_path
        except Exception as e2:
            # just copy overview as output
            import shutil
            shutil.copy(str(overview), str(output_path))
            return output_path

def export_docx(run_dir: pathlib.Path, output_path: pathlib.Path):
    """
    Export to DOCX via python-docx if available, else fallback to copying markdown.
    """
    overview = (run_dir / "overview.md")
    text = overview.read_text(encoding="utf-8") if overview.exists() else "No overview"
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        # narrow margins
        for s in doc.sections:
            s.top_margin = Inches(0.6)
            s.bottom_margin = Inches(0.6)
            s.left_margin = Inches(0.7)
            s.right_margin = Inches(0.7)

        # Title
        manifest={}
        try:
            import json
            manifest=json.loads((run_dir/"manifest.json").read_text(encoding="utf-8"))
        except Exception:
            pass
        t = doc.add_heading("SLM Evaluation Suite — Report", level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{manifest.get('model_id','')}  •  {manifest.get('preset','')}  •  {manifest.get('created_utc','')}")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(107,114,128)
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"Run ID {manifest.get('run_id','')}  •  Schema {manifest.get('schema_version','1.0')}")
        r2.font.size = Pt(7)
        r2.font.color.rgb = RGBColor(107,114,128)

        # body – simplistic markdown -> docx
        for line in text.split("\n"):
            stripped=line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("|"):
                # table row – add as monospace paragraph
                pp = doc.add_paragraph()
                rr = pp.add_run(stripped)
                rr.font.size = Pt(6)
                rr.font.name = "Courier New"
            elif stripped.startswith(">"):
                pp = doc.add_paragraph()
                pp.style = doc.styles['Quote'] if 'Quote' in [s.name for s in doc.styles] else pp.style
                rr = pp.add_run(stripped[1:].strip())
                rr.italic = True
                rr.font.size = Pt(8)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped.startswith("---"):
                doc.add_paragraph("_"*60)
            else:
                pp = doc.add_paragraph(stripped)
                pp.style.font.size = Pt(8)

        # charts
        doc.add_heading("Charts", level=1)
        charts = sorted((run_dir/"charts").glob("*.png"))
        for ch in charts:
            try:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rr = cap.add_run(ch.name)
                rr.font.size = Pt(7)
                rr.font.color.rgb = RGBColor(107,114,128)
                doc.add_picture(str(ch), width=Inches(6.0))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"Chart {ch.name} could not be embedded: {e}")

        doc.add_paragraph("")
        pp = doc.add_paragraph()
        rr = pp.add_run("Generated by slm-evaluation-suite v0.1.0 — evidence cards, not a single quality score. Descriptive only.")
        rr.font.size = Pt(7)
        rr.italic = True
        rr.font.color.rgb = RGBColor(107,114,128)

        doc.save(str(output_path))
        return output_path

    except ImportError:
        # fallback: save markdown as docx.txt
        fallback = output_path.with_suffix(".txt")
        fallback.write_text(text, encoding="utf-8")
        # also try copy
        if output_path.suffix != ".txt":
            try:
                import shutil
                shutil.copy(str(overview), str(output_path))
            except Exception:
                pass
        return fallback if fallback.exists() else output_path
